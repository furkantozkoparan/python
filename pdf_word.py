from docx import Document

# Belgenin yolunu belirtin
docx_path = 'C:/Users/furkan.tozkoparan/Desktop/icmal Test/temp.docx'

# Yeni ve eski isimleri belirtin
yeni_isim = "a"
eski_isim = "b"

# Word belgesini açın
doc = Document(docx_path)

# Belgedeki tüm paragrafları ve tabloları dönerek metni değiştirin
for paragraf in doc.paragraphs:
    for run in paragraf.runs:
        if eski_isim in run.text:
            run.text = run.text.replace(eski_isim, yeni_isim)
            run.font.highlight_color = None  # Vurguyu kaldır

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraf in cell.paragraphs:
                for run in paragraf.runs:
                    if eski_isim in run.text:
                        run.text = run.text.replace(eski_isim, yeni_isim)
                        run.font.highlight_color = None  # Vurguyu kaldır

# Değiştirilmiş ve vurgusuz Word belgesini kaydedin
updated_docx_path = 'C:/Users/furkan/Desktop/icmal Test/temp_updated.docx'
doc.save(updated_docx_path)
